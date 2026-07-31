"""In-image checks used by the release runtime regression harness.

This file deliberately generates all fixtures at runtime.  It is bind-mounted
into the image built from this repository, so the checks exercise the exact
LibreOffice, Python packages, Alembic migrations and ORM code that the
production image contains.  Tencent OCR's request contract is isolated behind
an in-process provider seam: this harness never needs a paid cloud credential.
It must never be pointed at a production database or uploads directory.
"""
from __future__ import annotations

import argparse
import hashlib
import os
import shutil
import subprocess
import sys
import tempfile
import time
from datetime import timedelta
from pathlib import Path
from uuid import uuid4


PROJECT_ROOT = Path("/app")
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise RuntimeError(message)


def _run_alembic(*arguments: str) -> None:
    """Run a migration without printing the database URL on failure."""

    completed = subprocess.run(
        ["alembic", *arguments],
        cwd=PROJECT_ROOT,
        text=True,
        capture_output=True,
        check=False,
    )
    if completed.returncode != 0:
        # Alembic normally does not print the configured connection URL, but
        # do not relay child output here because this harness is intentionally
        # safe to use with an ephemeral generated database password.
        raise RuntimeError(f"alembic_{arguments[0]}_failed")


def _make_pdf(path: Path, marker: str, *, font_size: float = 24) -> None:
    import fitz

    document = fitz.open()
    try:
        page = document.new_page(width=612, height=792)
        page.insert_text((54, 120), marker, fontsize=font_size, fontname="helv")
        document.save(path)
    finally:
        document.close()


def _make_image(path: Path, marker: str) -> None:
    """Render real raster text without adding a new image-library dependency."""

    import fitz

    with tempfile.TemporaryDirectory() as temporary:
        source_pdf = Path(temporary) / "image-source.pdf"
        _make_pdf(source_pdf, marker, font_size=34)
        document = fitz.open(source_pdf)
        try:
            pixmap = document.load_page(0).get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
            if path.suffix.lower() == ".png":
                pixmap.save(path)
            else:
                # PyMuPDF writes a real JPEG byte stream here. Giving this
                # path a .jpg suffix makes the application's actual image
                # route select Tencent OCR rather than a PDF fallback.
                path.write_bytes(pixmap.tobytes("jpeg"))
        finally:
            document.close()


def _assert_marker(result: object, *, parser_fragment: str, marker_words: tuple[str, ...]) -> None:
    parser_version = str(getattr(result, "parser_version"))
    raw_text = str(getattr(result, "raw_text"))
    _assert(parser_fragment in parser_version, f"unexpected_parser:{parser_version}")
    normalized = raw_text.upper()
    _assert(
        all(word in normalized for word in marker_words),
        f"synthetic_marker_missing:{parser_fragment}",
    )
    _assert(int(getattr(result, "source_page_count")) >= 1, "source_page_count_missing")
    _assert(int(getattr(result, "parsed_page_count")) >= 1, "parsed_page_count_missing")


def run_document_regression() -> None:
    """Exercise PDF, DOCX, XLSX, PNG, JPG and HTML in the production image."""

    from docx import Document
    from openpyxl import Workbook

    from app.services import document_text_extraction
    from app.services.tencent_ocr_provider import TencentOcrConfig

    _assert(
        shutil.which("tesseract") is None,
        "tesseract_must_not_be_present_in_production_image",
    )

    ocr_config = TencentOcrConfig(
        secret_id="runtime-regression-secret-id",
        secret_key="runtime-regression-secret-key",
        region="ap-guangzhou",
        timeout_seconds=5,
    )
    image_ocr_calls: list[str] = []

    def fake_image_ocr(*, path: Path, config: TencentOcrConfig) -> str:
        _assert(config == ocr_config, "unexpected_tencent_ocr_config")
        suffix = path.suffix.lower()
        _assert(suffix in {".png", ".jpg"}, "unexpected_tencent_image_suffix")
        image_ocr_calls.append(suffix)
        return f"SYNTHETIC {suffix[1:].upper()} RESUME MARKER"

    document_text_extraction.extract_image_text = fake_image_ocr
    extract_document_text = document_text_extraction.extract_document_text

    with tempfile.TemporaryDirectory(prefix="greatsell-document-regression-") as temporary:
        fixtures = Path(temporary)
        pdf_path = fixtures / "synthetic-resume.pdf"
        docx_path = fixtures / "synthetic-resume.docx"
        xlsx_path = fixtures / "synthetic-resume.xlsx"
        png_path = fixtures / "synthetic-resume.png"
        jpg_path = fixtures / "synthetic-resume.jpg"
        html_path = fixtures / "synthetic-resume.html"

        _make_pdf(pdf_path, "SYNTHETIC PDF RESUME MARKER")

        document = Document()
        document.add_heading("Synthetic resume", level=1)
        document.add_paragraph("SYNTHETIC DOCX RESUME MARKER")
        document.save(docx_path)

        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "Resume"
        worksheet.append(["SYNTHETIC", "XLSX", "RESUME", "MARKER"])
        workbook.save(xlsx_path)
        workbook.close()

        _make_image(png_path, "SYNTHETIC PNG RESUME MARKER")
        _make_image(jpg_path, "SYNTHETIC JPG RESUME MARKER")

        html_path.write_text(
            "<html><body><script>window.RUNTIME_SCRIPT_MARKER = true;</script>"
            "<main>SYNTHETIC HTML RESUME MARKER</main></body></html>",
            encoding="utf-8",
        )

        extraction_options = {
            "min_text_chars_per_page": 1,
            "ocr_sparse_text_chars_per_page": 1,
            "tencent_ocr_config": ocr_config,
        }
        pdf_result = extract_document_text(pdf_path, **extraction_options)
        _assert_marker(
            pdf_result,
            parser_fragment="pypdf-",
            marker_words=("SYNTHETIC", "PDF", "RESUME"),
        )

        docx_result = extract_document_text(docx_path, **extraction_options)
        _assert_marker(
            docx_result,
            parser_fragment="pypdf-",
            marker_words=("SYNTHETIC", "DOCX", "RESUME"),
        )

        xlsx_result = extract_document_text(xlsx_path, **extraction_options)
        _assert_marker(
            xlsx_result,
            parser_fragment="openpyxl",
            marker_words=("SYNTHETIC", "XLSX", "RESUME"),
        )

        png_result = extract_document_text(png_path, **extraction_options)
        _assert_marker(
            png_result,
            parser_fragment="tencent-ocr",
            marker_words=("SYNTHETIC", "PNG", "RESUME"),
        )

        jpg_result = extract_document_text(jpg_path, **extraction_options)
        _assert_marker(
            jpg_result,
            parser_fragment="tencent-ocr",
            marker_words=("SYNTHETIC", "JPG", "RESUME"),
        )

        html_result = extract_document_text(html_path, **extraction_options)
        _assert_marker(
            html_result,
            parser_fragment="beautifulsoup4",
            marker_words=("SYNTHETIC", "HTML", "RESUME"),
        )
        _assert(
            "RUNTIME_SCRIPT_MARKER" not in html_result.raw_text,
            "html_script_was_not_removed_before_extraction",
        )
        _assert(
            image_ocr_calls == [".png", ".jpg"],
            "tencent_image_ocr_route_was_not_selected",
        )

    print("runtime-document-regression: passed")


def _database_url() -> str:
    value = os.getenv("RESUME_V3_DATABASE_URL", "").strip()
    if not value.startswith("postgresql+"):
        raise RuntimeError("release_regression_requires_postgresql_url")
    return value


def _uploads_dir() -> Path:
    raw_path = os.getenv("RELEASE_REGRESSION_UPLOADS_DIR", "").strip()
    if not raw_path:
        raise RuntimeError("release_regression_uploads_dir_required")
    path = Path(raw_path)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _expected_alembic_head() -> str:
    from alembic.config import Config
    from alembic.script import ScriptDirectory

    return str(ScriptDirectory.from_config(Config(str(PROJECT_ROOT / "alembic.ini"))).get_current_head())


def _assert_current_head(database: object, expected_head: str) -> None:
    from sqlalchemy import text

    session_factory = getattr(database, "session_factory")
    with session_factory() as session:
        actual_head = session.scalar(text("SELECT version_num FROM alembic_version"))
    _assert(actual_head == expected_head, "alembic_head_mismatch")


def _settings(database_url: str, uploads_dir: Path) -> object:
    from app.config import AppSettings

    return AppSettings(
        project_dir=PROJECT_ROOT,
        data_dir=uploads_dir.parent,
        upload_dir=uploads_dir,
        database_url=database_url,
        auto_create_schema=False,
        seed_registry_on_startup=False,
        # This makes the extraction queue claimable without creating a
        # runnable legacy route or sending a provider request. The recovery
        # exercise stops after a lease-protected claim.
        ai_provider_credentials={"release-regression": "synthetic-not-a-secret"},
    )


def _seed_pre_0028_password_reset_rows(database_url: str) -> str:
    """Insert an account with the historical (pre-0028) table shape.

    ``PasswordResetToken`` in the current ORM knows about ``invalidated_at``,
    but that column deliberately does not exist at revision 0027.  Raw SQL is
    therefore intentional here: it proves that a populated database which
    issued multiple valid reset links before 0028 can really be upgraded.
    """

    from datetime import datetime, timezone

    from sqlalchemy import create_engine, text

    now = datetime.now(timezone.utc)
    user_id = str(uuid4())
    token_rows = []
    for position in (1, 2):
        token_rows.append(
            {
                "id": str(uuid4()),
                "user_id": user_id,
                "token_digest": hashlib.sha256(
                    f"release-regression-legacy-reset-{position}".encode("utf-8")
                ).hexdigest(),
                "expires_at": now + timedelta(hours=24),
                "requested_at": now - timedelta(minutes=position),
            }
        )

    engine = create_engine(database_url)
    try:
        with engine.begin() as connection:
            connection.execute(
                text(
                    """
                    INSERT INTO user_accounts (
                        id, email, email_key, full_name, password_hash,
                        is_active, is_platform_admin, email_verified_at,
                        last_login_at, created_at, updated_at
                    ) VALUES (
                        :id, :email, :email_key, :full_name, :password_hash,
                        :is_active, :is_platform_admin, :email_verified_at,
                        :last_login_at, :created_at, :updated_at
                    )
                    """
                ),
                {
                    "id": user_id,
                    "email": "release-regression-legacy-reset@invalid.test",
                    "email_key": "release-regression-legacy-reset@invalid.test",
                    "full_name": "Synthetic legacy reset account",
                    "password_hash": "!synthetic-release-regression-password!",
                    "is_active": True,
                    "is_platform_admin": False,
                    "email_verified_at": now,
                    "last_login_at": None,
                    "created_at": now,
                    "updated_at": now,
                },
            )
            # At 0027 this is the complete table shape: no ``invalidated_at``
            # column exists yet, and multiple unused rows are permitted.
            connection.execute(
                text(
                    """
                    INSERT INTO password_reset_tokens (
                        id, user_id, token_digest, expires_at, used_at, requested_at
                    ) VALUES (
                        :id, :user_id, :token_digest, :expires_at, NULL, :requested_at
                    )
                    """
                ),
                token_rows,
            )
    finally:
        engine.dispose()

    return user_id


def _assert_0028_password_reset_upgrade(database_url: str, *, user_id: str) -> None:
    """Prove 0028 invalidates legacy links and enforces one active link.

    This runs against PostgreSQL, rather than a mock or SQLite approximation:
    the second active insert must fail with the exact PostgreSQL partial-index
    name introduced by revision 0028.
    """

    from datetime import datetime, timezone

    from sqlalchemy import create_engine, text
    from sqlalchemy.exc import IntegrityError

    engine = create_engine(database_url)
    try:
        with engine.connect() as connection:
            legacy_rows = connection.execute(
                text(
                    """
                    SELECT id, used_at, invalidated_at
                    FROM password_reset_tokens
                    WHERE user_id = :user_id
                    ORDER BY requested_at ASC
                    """
                ),
                {"user_id": user_id},
            ).mappings().all()
            session_version = connection.scalar(
                text(
                    "SELECT auth_session_version FROM user_accounts WHERE id = :user_id"
                ),
                {"user_id": user_id},
            )

        _assert(len(legacy_rows) == 2, "legacy_password_reset_rows_missing")
        _assert(
            all(row["used_at"] is None for row in legacy_rows),
            "legacy_password_reset_rows_unexpectedly_used",
        )
        _assert(
            all(row["invalidated_at"] is not None for row in legacy_rows),
            "legacy_password_reset_rows_not_invalidated",
        )
        # Revision 0029 must also preserve the historical account while
        # adding its browser-session revocation counter.
        _assert(session_version == 1, "legacy_account_session_version_missing")

        now = datetime.now(timezone.utc)
        active_row = {
            "id": str(uuid4()),
            "user_id": user_id,
            "token_digest": hashlib.sha256(
                b"release-regression-post-upgrade-active-reset"
            ).hexdigest(),
            "expires_at": now + timedelta(hours=24),
            "requested_at": now,
        }
        with engine.begin() as connection:
            connection.execute(
                text(
                    """
                    INSERT INTO password_reset_tokens (
                        id, user_id, token_digest, expires_at, used_at,
                        invalidated_at, requested_at
                    ) VALUES (
                        :id, :user_id, :token_digest, :expires_at, NULL,
                        NULL, :requested_at
                    )
                    """
                ),
                active_row,
            )

        duplicate_row = {
            **active_row,
            "id": str(uuid4()),
            "token_digest": hashlib.sha256(
                b"release-regression-post-upgrade-duplicate-reset"
            ).hexdigest(),
            "requested_at": now + timedelta(seconds=1),
        }
        try:
            with engine.begin() as connection:
                connection.execute(
                    text(
                        """
                        INSERT INTO password_reset_tokens (
                            id, user_id, token_digest, expires_at, used_at,
                            invalidated_at, requested_at
                        ) VALUES (
                            :id, :user_id, :token_digest, :expires_at, NULL,
                            NULL, :requested_at
                        )
                        """
                    ),
                    duplicate_row,
                )
        except IntegrityError as exc:
            diagnostics = getattr(getattr(exc, "orig", None), "diag", None)
            constraint_name = getattr(diagnostics, "constraint_name", None)
            _assert(
                constraint_name == "uq_active_password_reset_per_user",
                "password_reset_partial_unique_index_not_enforced",
            )
        else:
            raise RuntimeError("password_reset_second_active_token_was_accepted")

        with engine.connect() as connection:
            active_count = connection.scalar(
                text(
                    """
                    SELECT count(*)
                    FROM password_reset_tokens
                    WHERE user_id = :user_id
                      AND used_at IS NULL
                      AND invalidated_at IS NULL
                    """
                ),
                {"user_id": user_id},
            )
        _assert(active_count == 1, "password_reset_active_token_count_mismatch")
    finally:
        engine.dispose()

    print("runtime-postgres-legacy-password-reset-migration: passed")


def run_database_seed() -> None:
    """Exercise both historical upgrade phases and write synthetic backup material."""

    from datetime import datetime, timezone

    from app.database import Database
    from app.models import (
        Candidate,
        MailboxBackgroundJob,
        MailboxConfig,
        Organization,
        Resume,
        ResumeAiExtractionJob,
        ResumeSourceBlock,
    )
    from app.tenant_scope import clear_organization_context, set_organization_context

    database_url = _database_url()
    uploads_dir = _uploads_dir()
    _run_alembic("upgrade", "20260716_0001")
    database = Database(database_url)
    try:
        _assert_current_head(database, "20260716_0001")
    finally:
        database.dispose()

    # Stage one explicitly proves the full historical chain through the last
    # revision before the password-reset security migration.  Stage two below
    # starts from a database that genuinely contains the old multi-link data.
    _run_alembic("upgrade", "20260721_0027")
    database = Database(database_url)
    try:
        _assert_current_head(database, "20260721_0027")
    finally:
        database.dispose()
    legacy_reset_user_id = _seed_pre_0028_password_reset_rows(database_url)

    _run_alembic("upgrade", "head")
    database = Database(database_url)
    expected_head = _expected_alembic_head()
    _assert_current_head(database, expected_head)
    _assert_0028_password_reset_upgrade(database_url, user_id=legacy_reset_user_id)

    payload = b"GreatSell synthetic recovery original. No candidate data.\n"
    digest = hashlib.sha256(payload).hexdigest()
    now = datetime.now(timezone.utc)
    try:
        with database.session_factory() as session:
            primary_organization = Organization(name="Synthetic recovery workspace")
            secondary_organization = Organization(name="Synthetic isolated workspace")
            session.add_all((primary_organization, secondary_organization))
            session.flush()

            set_organization_context(session, primary_organization.id)
            try:
                candidate = Candidate(display_name="Synthetic candidate")
                session.add(candidate)
                session.flush()
                storage_key = f"{primary_organization.id}/synthetic-recovery-resume.pdf"
                original_path = uploads_dir / storage_key
                original_path.parent.mkdir(parents=True, exist_ok=True)
                original_path.write_bytes(payload)
                resume = Resume(
                    candidate_id=candidate.id,
                    original_filename="synthetic-recovery-resume.pdf",
                    storage_key=storage_key,
                    sha256=digest,
                    source_page_count=1,
                    parsed_page_count=1,
                    extraction_status="text_ready",
                    quality_flags=[],
                    parser_version="release-regression",
                    is_active=False,
                    facts_version=0,
                    raw_text="SYNTHETIC RECOVERY MARKER",
                )
                mailbox = MailboxConfig(
                    display_name="Synthetic recovery mailbox",
                    display_name_key="synthetic recovery mailbox",
                    imap_host="imap.invalid.test",
                    imap_port=993,
                    email_address="synthetic-recovery@invalid.test",
                    mailbox="INBOX",
                    encrypted_password="synthetic-not-a-secret",
                    enabled=True,
                )
                session.add_all((resume, mailbox))
                session.flush()
                source_block = ResumeSourceBlock(
                    resume_id=resume.id,
                    block_id="page-001",
                    page_no=1,
                    block_type="text",
                    text="SYNTHETIC AI RECOVERY MARKER",
                )
                expired_mailbox_job = MailboxBackgroundJob(
                    mailbox_config_id=mailbox.id,
                    job_kind="sync",
                    trigger_type="manual",
                    status="running",
                    attempt_count=1,
                    max_attempts=3,
                    lease_owner="crashed-mailbox-worker",
                    lease_expires_at=now - timedelta(minutes=5),
                    requested_at=now - timedelta(minutes=10),
                    started_at=now - timedelta(minutes=6),
                )
                expired_ai_job = ResumeAiExtractionJob(
                    resume_id=resume.id,
                    job_kind="initial",
                    status="running",
                    attempt_count=1,
                    max_attempts=3,
                    input_facts_version=0,
                    lease_owner="crashed-ai-worker",
                    lease_expires_at=now - timedelta(minutes=5),
                    requested_at=now - timedelta(minutes=10),
                    started_at=now - timedelta(minutes=6),
                )
                session.add_all((source_block, expired_mailbox_job, expired_ai_job))
                # Flush before changing workspace context.  The ORM tenant
                # guard intentionally rejects a pending row from workspace A
                # if it is accidentally flushed while workspace B is active.
                session.flush()
            finally:
                clear_organization_context(session)

            set_organization_context(session, secondary_organization.id)
            try:
                secondary_candidate = Candidate(display_name="Synthetic isolated candidate")
                session.add(secondary_candidate)
                session.flush()
                secondary_resume = Resume(
                    candidate_id=secondary_candidate.id,
                    original_filename="synthetic-isolated-resume.pdf",
                    storage_key=f"{secondary_organization.id}/synthetic-isolated-resume.pdf",
                    sha256=hashlib.sha256(b"Synthetic isolated resume.").hexdigest(),
                    source_page_count=1,
                    parsed_page_count=1,
                    extraction_status="text_ready",
                    quality_flags=[],
                    parser_version="release-regression",
                    is_active=False,
                    facts_version=0,
                    raw_text="SYNTHETIC ISOLATED MARKER",
                )
                secondary_mailbox = MailboxConfig(
                    display_name="Synthetic isolated mailbox",
                    display_name_key="synthetic isolated mailbox",
                    imap_host="imap.invalid.test",
                    imap_port=993,
                    email_address="synthetic-isolated@invalid.test",
                    mailbox="INBOX",
                    encrypted_password="synthetic-not-a-secret",
                    enabled=True,
                )
                session.add_all((secondary_resume, secondary_mailbox))
                session.flush()
                untouched_mailbox_job = MailboxBackgroundJob(
                    mailbox_config_id=secondary_mailbox.id,
                    job_kind="sync",
                    trigger_type="manual",
                    status="queued",
                    attempt_count=0,
                    max_attempts=3,
                    next_attempt_at=now + timedelta(days=1),
                    requested_at=now,
                )
                untouched_ai_job = ResumeAiExtractionJob(
                    resume_id=secondary_resume.id,
                    job_kind="initial",
                    status="queued",
                    attempt_count=0,
                    max_attempts=3,
                    input_facts_version=0,
                    next_attempt_at=now + timedelta(days=1),
                    requested_at=now,
                )
                session.add_all((untouched_mailbox_job, untouched_ai_job))
                session.flush()
            finally:
                clear_organization_context(session)
            session.commit()
    finally:
        database.dispose()

    print("runtime-postgres-seed: passed")


def run_database_verify() -> None:
    """Verify restored data and run the worker's real expired-lease path."""

    from datetime import datetime, timezone

    from sqlalchemy import select

    from app.database import Database
    from app.models import MailboxBackgroundJob, Organization, Resume, ResumeAiExtractionJob
    from app.services import (
        ai_extraction_job_service,
        mailbox_background_job_service,
        workspace_background_lane_service,
    )
    from app.tenant_scope import clear_organization_context, set_organization_context

    database_url = _database_url()
    uploads_dir = _uploads_dir()
    database = Database(database_url)
    expected_head = _expected_alembic_head()
    _assert_current_head(database, expected_head)
    settings = _settings(database_url, uploads_dir)

    try:
        with database.session_factory() as session:
            primary = session.scalar(
                select(Organization).where(Organization.name == "Synthetic recovery workspace")
            )
            secondary = session.scalar(
                select(Organization).where(Organization.name == "Synthetic isolated workspace")
            )
            _assert(primary is not None and secondary is not None, "restored_workspaces_missing")

            set_organization_context(session, primary.id)
            try:
                resume = session.scalar(
                    select(Resume).where(
                        Resume.original_filename == "synthetic-recovery-resume.pdf"
                    )
                )
                expired_mailbox_job = session.scalar(
                    select(MailboxBackgroundJob).where(
                        MailboxBackgroundJob.status == "running",
                        MailboxBackgroundJob.lease_owner == "crashed-mailbox-worker",
                    )
                )
                expired_ai_job = session.scalar(
                    select(ResumeAiExtractionJob).where(
                        ResumeAiExtractionJob.status == "running",
                        ResumeAiExtractionJob.lease_owner == "crashed-ai-worker",
                    )
                )
            finally:
                clear_organization_context(session)
            _assert(resume is not None, "restored_resume_missing")
            _assert(expired_mailbox_job is not None, "restored_expired_mailbox_job_missing")
            _assert(expired_ai_job is not None, "restored_expired_ai_job_missing")
            _assert(expired_mailbox_job.organization_id == primary.id, "restored_mailbox_job_workspace_mismatch")
            _assert(expired_ai_job.organization_id == primary.id, "restored_ai_job_workspace_mismatch")
            original_path = uploads_dir / resume.storage_key
            _assert(original_path.is_file(), "restored_original_missing")
            _assert(
                hashlib.sha256(original_path.read_bytes()).hexdigest() == resume.sha256,
                "restored_original_sha256_mismatch",
            )

        # Invoke the exact mailbox worker recovery function, then its normal
        # claim function after the deliberate one-second retry backoff. No
        # IMAP connection is opened: this proves a restart can reclaim a
        # durable queue record before slow external work begins.
        with database.session_factory() as session:
            mailbox_background_job_service._recover_expired_jobs(
                session,
                settings=settings,
                now=datetime.now(timezone.utc),
            )

        with database.session_factory() as session:
            set_organization_context(session, primary.id)
            try:
                recovered_mailbox_job = session.get(MailboxBackgroundJob, expired_mailbox_job.id)
            finally:
                clear_organization_context(session)
            _assert(recovered_mailbox_job is not None, "recovered_mailbox_job_missing")
            _assert(recovered_mailbox_job.status == "queued", "expired_mailbox_lease_not_requeued")
            _assert(
                recovered_mailbox_job.last_error == "mailbox_background_job_lease_expired",
                "expired_mailbox_lease_error_missing",
            )
            _assert(
                recovered_mailbox_job.lease_owner is None
                and recovered_mailbox_job.lease_expires_at is None,
                "expired_mailbox_lease_not_cleared",
            )
            _assert(
                recovered_mailbox_job.next_attempt_at is not None,
                "expired_mailbox_lease_retry_not_scheduled",
            )
            wait_seconds = max(
                0.0,
                (
                    recovered_mailbox_job.next_attempt_at
                    - datetime.now(timezone.utc)
                ).total_seconds(),
            )

        if wait_seconds:
            time.sleep(wait_seconds + 0.15)
        claimed = mailbox_background_job_service._claim_next_job(
            database,
            settings=settings,
            worker_id="recovery-regression-worker",
        )
        _assert(claimed is not None, "recovered_mailbox_job_not_claimable")
        _assert(claimed.organization_id == primary.id, "recovered_mailbox_job_claimed_cross_workspace")

        # The direct claim above intentionally stops before the mailbox worker
        # performs an IMAP request.  It must nevertheless reserve the shared
        # heavy-work lane, so an AI job in the same workspace cannot run at the
        # same time.
        blocked_ai_claim = ai_extraction_job_service._claim_next_job(
            database,
            settings=settings,
            worker_id="recovery-regression-blocked-ai-worker",
        )
        _assert(
            blocked_ai_claim is None,
            "same_workspace_mailbox_lane_did_not_block_ai_claim",
        )

        # Simulate the normal mailbox worker completion/finally sequence
        # without opening an IMAP connection: complete the claimed job, then
        # release the exact fenced lane token.  The following AI claim proves
        # that recovery does not leave capacity stranded after completion.
        with database.session_factory() as session:
            set_organization_context(session, claimed.organization_id)
            try:
                mailbox_completed = mailbox_background_job_service._complete_job(
                    session,
                    claimed=claimed,
                    worker_id="recovery-regression-worker",
                )
            finally:
                clear_organization_context(session)
            _assert(mailbox_completed, "recovered_mailbox_job_not_completed")
            lane_released = workspace_background_lane_service.release_workspace_background_lane(
                session,
                organization_id=claimed.organization_id,
                lease_token=claimed.workspace_lane_token,
            )
            session.commit()
        _assert(lane_released, "recovered_mailbox_lane_not_released")

        # AI extraction has a different lease/retry implementation from the
        # mailbox queue. Exercise its own recovery and global claim code with
        # an inert in-memory credential map; no provider execution follows.
        with database.session_factory() as session:
            ai_extraction_job_service._recover_expired_leases(
                session,
                now=datetime.now(timezone.utc),
            )
            session.commit()

        with database.session_factory() as session:
            set_organization_context(session, primary.id)
            try:
                recovered_ai_job = session.get(ResumeAiExtractionJob, expired_ai_job.id)
            finally:
                clear_organization_context(session)
            _assert(recovered_ai_job is not None, "recovered_ai_job_missing")
            _assert(recovered_ai_job.status == "queued", "expired_ai_lease_not_requeued")
            _assert(
                recovered_ai_job.last_error == "ai_extraction_worker_lease_expired",
                "expired_ai_lease_error_missing",
            )
            _assert(
                recovered_ai_job.lease_owner is None
                and recovered_ai_job.lease_expires_at is None,
                "expired_ai_lease_not_cleared",
            )

        ai_claimed = ai_extraction_job_service._claim_next_job(
            database,
            settings=settings,
            worker_id="recovery-regression-ai-worker",
        )
        _assert(ai_claimed is not None, "recovered_ai_job_not_claimable")
        _assert(ai_claimed.organization_id == primary.id, "recovered_ai_job_claimed_cross_workspace")

        with database.session_factory() as session:
            set_organization_context(session, primary.id)
            try:
                completed_mailbox_job = session.get(MailboxBackgroundJob, expired_mailbox_job.id)
                reclaimed_ai_job = session.get(ResumeAiExtractionJob, expired_ai_job.id)
            finally:
                clear_organization_context(session)
            set_organization_context(session, secondary.id)
            try:
                untouched_mailbox_job = session.scalar(
                    select(MailboxBackgroundJob).where(
                        MailboxBackgroundJob.status == "queued"
                    )
                )
                untouched_ai_job = session.scalar(
                    select(ResumeAiExtractionJob).where(
                        ResumeAiExtractionJob.status == "queued"
                    )
                )
            finally:
                clear_organization_context(session)
            _assert(completed_mailbox_job is not None, "completed_mailbox_job_missing")
            _assert(completed_mailbox_job.status == "completed", "recovered_mailbox_job_not_completed")
            _assert(
                completed_mailbox_job.lease_owner is None,
                "completed_mailbox_job_lease_owner_not_cleared",
            )
            _assert(reclaimed_ai_job is not None, "reclaimed_ai_job_missing")
            _assert(reclaimed_ai_job.status == "running", "recovered_ai_job_not_running_after_claim")
            _assert(
                reclaimed_ai_job.lease_owner == "recovery-regression-ai-worker",
                "recovered_ai_job_owner_mismatch",
            )
            _assert(untouched_mailbox_job is not None, "secondary_workspace_mailbox_job_missing")
            _assert(untouched_mailbox_job.organization_id == secondary.id, "secondary_workspace_mailbox_job_changed")
            _assert(untouched_mailbox_job.attempt_count == 0, "secondary_workspace_mailbox_job_claimed")
            _assert(untouched_ai_job is not None, "secondary_workspace_ai_job_missing")
            _assert(untouched_ai_job.organization_id == secondary.id, "secondary_workspace_ai_job_changed")
            _assert(untouched_ai_job.attempt_count == 0, "secondary_workspace_ai_job_claimed")
    finally:
        database.dispose()

    print("runtime-postgres-restore-and-lease-recovery: passed")


def main() -> None:
    parser = argparse.ArgumentParser(description="Run release regression checks inside the application image.")
    parser.add_argument(
        "mode",
        choices=("documents", "database-seed", "database-verify"),
        help="The isolated runtime check to execute.",
    )
    arguments = parser.parse_args()
    if arguments.mode == "documents":
        run_document_regression()
    elif arguments.mode == "database-seed":
        run_database_seed()
    else:
        run_database_verify()


if __name__ == "__main__":
    main()
